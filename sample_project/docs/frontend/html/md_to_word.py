import os
import re
import requests
import base64
from docx import Document
from docx.shared import Inches
from io import BytesIO

def render_mermaid(mermaid_code):
    """Converts Mermaid syntax to an image via mermaid.ink API."""
    # Encode Mermaid code to base64
    graph_bytes = mermaid_code.encode("utf-8")
    base64_bytes = base64.b64encode(graph_bytes)
    base64_string = base64_bytes.decode("utf-8")
    
    # Use mermaid.ink to get a PNG
    url = f"https://mermaid.ink/img/{base64_string}"
    response = requests.get(url)
    
    if response.status_code == 200:
        return BytesIO(response.content)
    else:
        print(f"Failed to render diagram. Status: {response.status_code}")
        return None

def convert_md_to_word(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    doc.add_heading('Documentation Report', 0)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split content by mermaid blocks
    # This regex finds ```mermaid ... ```
    chunks = re.split(r'```mermaid\n(.*?)\n```', content, flags=re.DOTALL)

    for i, chunk in enumerate(chunks):
        if i % 2 == 1:
            # This is a mermaid block
            print(f"Rendering diagram {i//2 + 1}...")
            img_data = render_mermaid(chunk.strip())
            if img_data:
                doc.add_picture(img_data, width=Inches(6))
                doc.add_paragraph("[Diagram Rendered Above]")
        else:
            # This is regular text/markdown (Simple parsing)
            lines = chunk.strip().split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip() == '---':
                    doc.add_page_break()
                elif line.strip():
                    doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Successfully created: {docx_path}")

if __name__ == "__main__":
    # Update these paths as needed
    input_md = r'c:\Users\pavan\OneDrive\Desktop\Python FSD\django\calculator_project\docs\frontend\html\readme.md'
    output_docx = r'c:\Users\pavan\OneDrive\Desktop\Python FSD\django\calculator_project\docs\frontend\html\Documentation.docx'
    
    convert_md_to_word(input_md, output_docx)